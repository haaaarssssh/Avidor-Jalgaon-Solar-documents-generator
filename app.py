from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from docx import Document
from docx.shared import Inches
import os
import zipfile
from io import BytesIO
import tempfile
import traceback
from werkzeug.utils import secure_filename
from datetime import datetime
from PIL import Image
import base64
from io import BytesIO
import subprocess
import re

app = Flask(__name__)
app.secret_key = 'solar_unified_doc_generator_2025_secure_key'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

VARIABLE_MAPPING = {
    'name_variable': 'consumer_name',
    'consumer_number_variable': 'consumer_number',
    'consumer_variable': 'consumer_number',
    'address_variable': 'address',
    'sanctioned_capacity_variable': 'sanctioned_capacity',
    'reinstalled_capacity_variable': 'installed_capacity',
    'module_make_variable': 'module_make',
    'inverter_capacity_variable': 'inverter_capacity',
    'module_capacity_variable': 'module_capacity',
    'number_of_pv_modules_variable': 'number_of_modules',
    'district_variable': 'district',
    'installation_date_variable': 'installation_date',
    'distribution_license_variable': 'distribution_licensee',
    'model_number_variable': 'model_number',
    'wattage_variable': 'wattage',
    'model_number_inverter_variable': 'model_number_inverter',
    'rating_variable': 'rating',
    'aadhar_number_variable': 'aadhar_number',
    'executed_date_variable': 'agreement_date',
    'module_number': 'model_number',
    'model_capacity': 'model_capacity',
    'sanctioned_caacity_variable': 'sanctioned_capacity',
    'cost_of_rts_variable': 'total_cost',
    'mobile_number_variable': 'mobile_number',
    'email_variable': 'email',
    'system_checkdate_variable': 'performance_check_date',
    'todays_date_variable': 'todays_date',
}

DOCUMENT_TEMPLATES = {
    'NET': 'static/templates/NET.docx',
    'WCR': 'static/templates/WCR.docx',
    'Model-Agreement': 'static/templates/Model-Agreement.docx',
    'Proforma-A': 'static/templates/2.-Annexure-I-Profarma-A.docx'
}

def check_libreoffice():
    """Check if LibreOffice is available"""
    try:
        result = subprocess.run(['libreoffice', '--version'], 
                              capture_output=True, timeout=5)
        return result.returncode == 0
    except:
        return False

PDF_AVAILABLE = check_libreoffice()

def replace_in_runs(runs, replacements):
    """Handle variables split across runs (CASE-INSENSITIVE) - Makes ONLY replaced values bold"""
    if not runs:
        return

    full_text = ''.join(run.text for run in runs)

    # Track segments: list of (text, is_replacement)
    segments = []
    remaining_text = full_text
    modified = False

    # Create case-insensitive pattern for all variables
    import re

    # Sort variables by length (longest first) to avoid partial replacements
    sorted_vars = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)

    # Find all variable positions and their replacements
    replacements_map = []  # List of (start, end, replacement_value)

    for var_name, var_value in sorted_vars:
        # Create case-insensitive pattern
        pattern = re.compile(re.escape(var_name), re.IGNORECASE)

        # Find all matches
        for match in pattern.finditer(full_text):
            start, end = match.span()
            # Check if this position is already covered
            overlap = False
            for existing_start, existing_end, _ in replacements_map:
                if not (end <= existing_start or start >= existing_end):
                    overlap = True
                    break

            if not overlap:
                replacements_map.append((start, end, str(var_value)))
                modified = True

    if not modified:
        return

    # Sort replacements by position
    replacements_map.sort(key=lambda x: x[0])

    # Build segments
    current_pos = 0
    for start, end, replacement_value in replacements_map:
        # Add text before replacement (if any)
        if start > current_pos:
            segments.append((full_text[current_pos:start], False))

        # Add replacement value
        segments.append((replacement_value, True))

        current_pos = end

    # Add remaining text
    if current_pos < len(full_text):
        segments.append((full_text[current_pos:], False))

    # Get the original formatting from the first run
    original_font_name = runs[0].font.name if runs[0].font.name else None
    original_font_size = runs[0].font.size if runs[0].font.size else None

    # Clear all existing runs except the first
    for run in runs[1:]:
        run.text = ''

    # Clear the first run
    runs[0].text = ''

    # Get the paragraph to add new runs
    paragraph = runs[0]._element.getparent()

    # Remove all run elements
    for run in runs:
        if run._element.getparent() is not None:
            run._element.getparent().remove(run._element)

    # Add new runs for each segment
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for text, is_replacement in segments:
        if text:  # Only add non-empty segments
            # Create new run element
            new_run_elem = OxmlElement('w:r')
            paragraph.append(new_run_elem)

            # Create run properties
            rPr = OxmlElement('w:rPr')
            new_run_elem.append(rPr)

            # Set font properties
            if original_font_name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), original_font_name)
                rPr.append(rFonts)

            if original_font_size:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(original_font_size.pt * 2)))  # Word uses half-points
                rPr.append(sz)

            # Make bold ONLY if it's a replacement
            if is_replacement:
                b = OxmlElement('w:b')
                rPr.append(b)

            # Add text
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            # Preserve spaces
            text_elem.set(qn('xml:space'), 'preserve')
            new_run_elem.append(text_elem)



def docx_replace_robust(doc, form_data):
    """Replace variables in paragraphs AND tables"""
    replacements = {}
    for variable, field_name in VARIABLE_MAPPING.items():
        if field_name in form_data and form_data[field_name]:
            replacements[variable] = form_data[field_name]
    
    # Process paragraphs
    for para in doc.paragraphs:
        replace_in_runs(para.runs, replacements)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, replacements)
                
                # Process nested tables
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for para in nested_cell.paragraphs:
                                replace_in_runs(para.runs, replacements)
    
    # Process headers/footers
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                replace_in_runs(para.runs, replacements)
        except:
            pass
        try:
            for para in section.footer.paragraphs:
                replace_in_runs(para.runs, replacements)
        except:
            pass

def add_images_to_wcr(doc, aadhar_path, signature_path):
    """Add consumer signature and aadhar images"""
    try:
        replaced_count = 0
        
        for para in doc.paragraphs:
            para_text = para.text
            
            # Consumer signature - replace the text even if vendor image exists
            if 'signature_image_variable' in para_text:
                for run in para.runs:
                    if 'signature_image_variable' in run.text:
                        run.text = run.text.replace('signature_image_variable', '')
                        if signature_path and os.path.exists(signature_path):
                            run.add_picture(signature_path, width=Inches(1.0))
                            replaced_count += 1
                            print(f"  ✓ Added consumer signature #{replaced_count}")
            
            # Aadhar image
            if 'aadhar_image_variable' in para_text:
                for run in para.runs:
                    if 'aadhar_image_variable' in run.text:
                        run.text = run.text.replace('aadhar_image_variable', '')
                        if aadhar_path and os.path.exists(aadhar_path):
                            run.add_picture(aadhar_path, width=Inches(2.0))
                            print(f"  ✓ Added Aadhar image")
        
        if replaced_count > 0:
            print(f"  ✓ Total signature images added: {replaced_count}")
        
        return True
        
    except Exception as e:
        print(f"  ✗ Image error: {e}")
        traceback.print_exc()
        return True


def add_signature_to_proforma(doc, signature_path):
    """Replace signature_image_variable text with actual signature image"""
    if not signature_path or not os.path.exists(signature_path):
        return False
    
    try:
        replaced_count = 0
        
        # Process all paragraphs
        for para in doc.paragraphs:
            if 'signature_image_variable' in para.text:
                # Replace the text with image (SMALLER SIZE)
                for run in para.runs:
                    if 'signature_image_variable' in run.text:
                        run.text = run.text.replace('signature_image_variable', '')
                        # Changed from 1.5 to 1.0 inches
                        run.add_picture(signature_path, width=Inches(0.8))
                        replaced_count += 1
                        print(f"  ✓ Added signature image #{replaced_count}")
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if 'signature_image_variable' in para.text:
                            for run in para.runs:
                                if 'signature_image_variable' in run.text:
                                    run.text = run.text.replace('signature_image_variable', '')
                                    # Changed from 1.5 to 1.0 inches
                                    run.add_picture(signature_path, width=Inches(1.0))
                                    replaced_count += 1
                                    print(f"  ✓ Added signature image in table #{replaced_count}")
        
        if replaced_count > 0:
            print(f"  ✓ Total signature images added: {replaced_count}")
            return True
        else:
            print(f"  ⚠️ No signature_image_variable found")
            return False
        
    except Exception as e:
        print(f"  ✗ Error adding signature: {e}")
        traceback.print_exc()
        return False

def process_cropped_image(base64_data):
    """Convert base64 cropped image to file"""
    try:
        if ',' in base64_data:
            base64_data = base64_data.split(',')[1]
        
        image_data = base64.b64decode(base64_data)
        image = Image.open(BytesIO(image_data))
        
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f'cropped_{datetime.now().timestamp()}.jpg')
        image.save(temp_path, 'JPEG', quality=90)
        
        return temp_path
    except Exception as e:
        print(f"Error processing cropped image: {e}")
        return None


def convert_to_pdf_libreoffice(docx_path):
    """Convert DOCX to PDF using LibreOffice with better formatting"""
    try:
        if not PDF_AVAILABLE:
            print(f"  ⚠️ LibreOffice not available")
            return None
        
        output_dir = os.path.dirname(docx_path)
        
        # Use better conversion options
        result = subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf:writer_pdf_Export',  # Explicit PDF filter
            '--outdir', output_dir,
            docx_path
        ], capture_output=True, timeout=90, text=True)  # Increased timeout
        
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            print(f"  ✓ PDF created ({file_size:,} bytes)")
            return pdf_path
        else:
            print(f"  ✗ PDF creation failed")
            if result.stderr:
                print(f"  Error: {result.stderr}")
            return None
            
    except subprocess.TimeoutExpired:
        print(f"  ✗ PDF conversion timeout (90s)")
        return None
    except Exception as e:
        print(f"  ✗ PDF error: {e}")
        traceback.print_exc()
        return None


@app.route('/')
def index():
    return render_template('unified_form.html', pdf_available=PDF_AVAILABLE)

@app.route('/generate_documents', methods=['POST'])
def generate_documents():
    print("\n" + "="*80)
    print("STARTING DOCUMENT GENERATION")
    print(f"LibreOffice available: {PDF_AVAILABLE}")
    print("="*80 + "\n")
    
    try:
        # Get format option
        output_format = request.form.get('output_format', 'both')
        print(f"Output format: {output_format}")
        
        # Collect form data
        form_data = {}
        for key in request.form:
            value = request.form.get(key, '').strip()
            form_data[key] = value
        
        form_data['todays_date'] = datetime.now().strftime('%d/%m/%Y')
        print(f"✓ Collected {len(form_data)} form fields")
        
        # Handle file uploads
        uploaded_files = {}
        signature_cropped = request.form.get('signature_image_cropped')
        if signature_cropped:
            sig_path = process_cropped_image(signature_cropped)
            if sig_path:
                uploaded_files['signature_image'] = sig_path
                print(f"✓ Signature cropped and saved")

        aadhar_cropped = request.form.get('aadhar_image_cropped')
        if aadhar_cropped:
            aadhar_path = process_cropped_image(aadhar_cropped)
            if aadhar_path:
                uploaded_files['aadhar_image'] = aadhar_path
                print(f"✓ Aadhar cropped and saved")
        
        # Create temp directory
        tmpdir = tempfile.mkdtemp()
        print(f"✓ Temp dir: {tmpdir}\n")
        
        generated_files = []
        errors = []
        
        # Process each document
        for doc_name, template_path in DOCUMENT_TEMPLATES.items():
            print(f"{'='*60}")
            print(f"Processing: {doc_name}")
            print(f"{'='*60}")
            
            if not os.path.exists(template_path):
                error_msg = f"Template not found: {template_path}"
                print(f"✗ {error_msg}")
                errors.append(error_msg)
                continue
            
            try:
                # Load document
                doc = Document(template_path)
                print(f"  ✓ Loaded ({len(doc.paragraphs)} paras, {len(doc.tables)} tables)")
                
                # Replace variables
                docx_replace_robust(doc, form_data)
                print(f"  ✓ Variables replaced")
                
                # Add images to WCR
                if doc_name == 'WCR':
                    aadhar = uploaded_files.get('aadhar_image')
                    sig = uploaded_files.get('signature_image')
                    if aadhar or sig:
                        add_images_to_wcr(doc, aadhar, sig)

                # Add signature to Proforma-A
                if doc_name == 'Proforma-A':
                    sig = uploaded_files.get('signature_image')
                    if sig:
                        add_signature_to_proforma(doc, sig)

                # Save DOCX (always need to save first)
                docx_file = os.path.join(tmpdir, f"{doc_name}.docx")
                doc.save(docx_file)
                
                if not os.path.exists(docx_file):
                    error_msg = f"{doc_name}: DOCX save failed"
                    print(f"  ✗ {error_msg}")
                    errors.append(error_msg)
                    continue
                
                print(f"  ✓ DOCX saved ({os.path.getsize(docx_file):,} bytes)")
                
                # Handle format selection
                if output_format == 'docx':
                    # User wants DOCX only
                    generated_files.append(docx_file)
                    
                elif output_format == 'pdf':
                    # User wants PDF only
                    if not PDF_AVAILABLE:
                        error_msg = f"{doc_name}: PDF not available (LibreOffice not installed)"
                        print(f"  ✗ {error_msg}")
                        errors.append(error_msg)
                        continue
                    
                    print(f"  → Converting to PDF...")
                    pdf_file = convert_to_pdf_libreoffice(docx_file)
                    
                    if pdf_file:
                        generated_files.append(pdf_file)
                        # Delete DOCX since user only wants PDF
                        try:
                            os.remove(docx_file)
                        except:
                            pass
                    else:
                        error_msg = f"{doc_name}: PDF conversion failed"
                        print(f"  ✗ {error_msg}")
                        errors.append(error_msg)
                
                elif output_format == 'both':
                    # User wants both DOCX and PDF
                    generated_files.append(docx_file)
                    
                    if PDF_AVAILABLE:
                        print(f"  → Converting to PDF...")
                        pdf_file = convert_to_pdf_libreoffice(docx_file)
                        if pdf_file:
                            generated_files.append(pdf_file)
                        else:
                            print(f"  ⚠️ PDF conversion failed, but DOCX available")
                    else:
                        print(f"  ⚠️ PDF not available (LibreOffice not installed)")
                
            except Exception as e:
                error_msg = f"{doc_name}: {str(e)}"
                print(f"  ✗ ERROR: {error_msg}")
                traceback.print_exc()
                errors.append(error_msg)
            
            print()
        
        print(f"{'='*60}")
        print(f"Generated {len(generated_files)} files")
        if errors:
            print(f"Errors: {len(errors)}")
            for err in errors:
                print(f"  - {err}")
        print(f"{'='*60}\n")
        
        if not generated_files:
            if output_format == 'pdf' and not PDF_AVAILABLE:
                flash('PDF generation not available. Please install LibreOffice or select DOCX format.', 'error')
            else:
                flash('No documents were generated! Check console for errors.', 'error')
            return redirect(url_for('index'))
        
        # Create ZIP
        consumer_name = form_data.get('consumer_name', 'client').replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f"Solar_Documents_{consumer_name}_{timestamp}.zip"
        zip_path = os.path.join(tmpdir, zip_filename)
        
        print(f"Creating ZIP: {zip_filename}")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in generated_files:
                arcname = os.path.basename(file_path)
                zipf.write(file_path, arcname=arcname)
                print(f"  ✓ {arcname}")
        
        # Read ZIP into memory
        with open(zip_path, 'rb') as f:
            zip_data = BytesIO(f.read())
        zip_data.seek(0)
        
        # Cleanup uploaded files
        for file_path in uploaded_files.values():
            try:
                os.remove(file_path)
            except:
                pass
        
        print(f"\n✓ Sending ZIP\n" + "="*80 + "\n")
        
        return send_file(
            zip_data,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
    
    except Exception as e:
        print(f"\nFATAL ERROR: {e}")
        traceback.print_exc()
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
